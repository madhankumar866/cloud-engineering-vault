from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
# set landscape
section = doc.sections[-1]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

# Set small margins
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# Add Title
p = doc.add_paragraph("Terraform - Tracker")
p.alignment = 1 # Center

table = doc.add_table(rows=1, cols=32)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Topics / Days'
for i in range(1, 32):
    hdr_cells[i].text = str(i)
    hdr_cells[i].width = Inches(0.2)

topics = [
    "🏗️ Terraform Core Building Blocks",
    "HCL Fundamentals",
    "State Management",
    "Modules",
    "⚙️ Advanced Configuration",
    "Loops & Meta-Arguments",
    "Workspaces & Environments",
    "Testing & Validation",
    "☁️ Ecosystem & Operations",
    "EKS & GitOps Bootstrapping",
    "FinOps & Cost Optimization (Infracost)",
    "⏳ Pending Topics",
    "Terraform Cloud / Enterprise",
    "Custom Providers",
    "🛠️ Hands-On Labs",
    "Lab 1: State & Modules",
    "Lab 2: Loops & Workspaces",
    "Lab 3: 3-Tier Capstone",
    "Lab 4: EKS & GitOps",
    "🏃 Habits",
    "Gym",
    "Book Reading",
    "Swimming"
]

for topic in topics:
    row_cells = table.add_row().cells
    row_cells[0].text = topic
    
    # Shade header rows
    if topic.startswith("🏗️") or topic.startswith("⚙️") or topic.startswith("☁️") or topic.startswith("⏳") or topic.startswith("🛠️") or topic.startswith("🏃"):
        # Make bold
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        # Add background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'EAEAEA')
        for cell in row_cells:
            cell._tc.get_or_add_tcPr().append(shading_elm)

doc.save('/Users/mk/Documents/01_Active_Work/obsidian/Cloud-Engineering-Vault/Trackers/Terraform_Habit_Tracker_Landscape.docx')
