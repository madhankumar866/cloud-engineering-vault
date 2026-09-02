from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Helvetica'
font.size = Pt(8)

# set landscape
section = doc.sections[-1]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

# Set very small margins
section.top_margin = Inches(0.2)
section.bottom_margin = Inches(0.2)
section.left_margin = Inches(0.2)
section.right_margin = Inches(0.2)

# Add Title
p = doc.add_paragraph("Terraform - Tracker")
p.alignment = 1 # Center
p.runs[0].font.size = Pt(12)
p.runs[0].font.bold = True

table = doc.add_table(rows=1, cols=32)
table.style = 'Table Grid'
# Ensure auto-fit is off so columns stay tight
table.autofit = False

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Topics / Days'
hdr_cells[0].width = Inches(2.0)
for i in range(1, 32):
    hdr_cells[i].text = str(i)
    hdr_cells[i].width = Inches(0.25)

topics = [
    "🏗️ Terraform Core Building Blocks",
    "HCL Fundamentals",
    "State Management",
    "Modules",
    "⚙️ Advanced Configuration",
    "Loops & Meta-Arguments",
    "Workspaces & Environments",
    "Testing & Validation",
    "☁️ Ecosystem & Operations",
    "EKS & GitOps Bootstrapping",
    "FinOps & Cost Optimization",
    "⏳ Pending Topics",
    "Terraform Cloud / Enterprise",
    "Custom Providers",
    "🛠️ Hands-On Labs",
    "Lab 1: State & Modules",
    "Lab 2: Loops & Workspaces",
    "Lab 3: 3-Tier Capstone",
    "Lab 4: EKS & GitOps",
    "🏃 Habits",
    "Gym",
    "Book Reading",
    "Swimming"
]

for topic in topics:
    row = table.add_row()
    row.height = Inches(0.2) # Tighten row height
    row_cells = row.cells
    row_cells[0].text = topic
    row_cells[0].width = Inches(2.0)
    for i in range(1, 32):
        row_cells[i].width = Inches(0.25)
    
    # Shade header rows
    if topic.startswith("🏗️") or topic.startswith("⚙️") or topic.startswith("☁️") or topic.startswith("⏳") or topic.startswith("🛠️") or topic.startswith("🏃"):
        # Make bold
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        # Add background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'EAEAEA')
        for cell in row_cells:
            cell._tc.get_or_add_tcPr().append(shading_elm)

# Apply font size 8 to all cells explicitly just in case
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

doc.save('/Users/mk/Documents/01_Active_Work/obsidian/Cloud-Engineering-Vault/Trackers/Terraform_Habit_Tracker_Landscape.docx')
